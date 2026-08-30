"""Monta el Word con todas las pantallas de la aplicación.

Una pantalla por página: el pantallazo de escritorio grande y, al lado, el del
móvil en pequeño. La idea es poder enseñar la aplicación sin abrirla, así que
cada una lleva una frase de qué es y no un pie de foto que repita el título.

Uso: python scripts/documento.py [carpeta de capturas] [fichero de salida]
"""

import os
import struct
import sys
from datetime import date

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

CAPTURAS = sys.argv[1] if len(sys.argv) > 1 else 'capturas/doc'
# Lo que cabe de alto en un A4 quitando márgenes, título y descripción.
ALTO_UTIL = 20.5
SALIDA = sys.argv[2] if len(sys.argv) > 2 else 'Gastos - pantallas.docx'

# Los mismos colores de la aplicación, para que el documento no desentone.
TINTA = RGBColor(0x11, 0x11, 0x11)
TINTA_2 = RGBColor(0x6B, 0x6B, 0x66)

# El orden es el del menú: así se lee como se usa.
PANTALLAS = [
    ('mes', 'Mes', 'La pantalla de cada día. Arriba, una sola cifra: lo que te queda. '
     'Debajo, tres bloques que responden a las tres preguntas de después —cuánto se va '
     'solo, cómo va el sobre de la comida y en qué se va lo demás— y las dos listas con '
     'las que se trabaja.'),
    ('mes-menu', 'Mes · acciones del mes',
     'Lo que no es del día a día: regenerar desde la plantilla, reiniciar, borrar y '
     'cerrar el mes, más el objetivo de ahorro de este mes. Cuando algo no se puede '
     'deshacer, la confirmación sustituye a la lista aquí mismo y dice con números qué '
     'se pierde.'),
    ('mes-analisis', 'Mes · análisis',
     'Plegado dentro de Mes, porque se mira de vez en cuando y no cada día: en qué se '
     'ha ido el dinero, la regla 50/30/20 contra tus propios ideales y los conceptos '
     'que más pesan.'),
    ('anual', 'Año',
     'La matriz de concepto por mes. La primera columna se queda fija al desplazar de '
     'lado, cada fila lleva su línea de evolución y cada celda lleva a su mes. A la '
     'derecha, el total, la media y la comparación con el año anterior.'),
    ('analitica-evolucion', 'Analítica · Evolución',
     'Cómo va cambiando lo que eliges, mes a mes, en el rango que elijas.'),
    ('analitica-anios', 'Analítica · Años',
     'El mismo año contra otro, superpuestos, para ver si vas mejor o peor.'),
    ('analitica-reparto', 'Analítica · Reparto',
     'En qué se reparte el dinero y cómo cambia esa proporción con el tiempo.'),
    ('analitica-ahorro', 'Analítica · Ahorro',
     'Cuánto se aparta de verdad cada mes, y contra qué objetivo.'),
    ('analitica-meses', 'Analítica · Meses',
     'Qué se dispara en un mes concreto. El mapa de calor compara cada concepto '
     'consigo mismo, no con los demás: lo que se busca es la forma.'),
    ('conceptos', 'Conceptos',
     'El catálogo. Cada concepto tiene su color y su icono, que se eligen solos por el '
     'nombre y se pueden cambiar tocando el icono de la fila. Se ordenan arrastrando.'),
    ('conceptos-plantilla', 'Conceptos · Plantilla',
     'Lo que costará un mes antes de que pase nada. Los importes tienen histórico: lo '
     'que se cambia vale desde el mes que elijas arriba y lo anterior no se toca.'),
    ('importar-extracto', 'Importar · Extracto del banco',
     'Se arrastra el archivo del banco y la aplicación lo reparte: concilia los fijos, '
     'crea los variables y la comida, y deja solo lo que no reconoce. Nada se guarda '
     'hasta que se acepta, y se puede deshacer entero.'),
    ('revision', 'Importar · Revisión del extracto',
     'La revisión antes de aceptar. Arriba el recuento; lo que falta por clasificar va '
     'en ámbar y con su propio bloque. Debajo de cada descripción limpia está la del '
     'banco, que es la que se coteja.'),
    ('importar-excel', 'Importar · Excel histórico',
     'Para traer los años que estaban en la hoja de cálculo.'),
    ('importar-copia', 'Importar · Copia de seguridad',
     'Exportar todo a JSON o a Excel, y volver a meterlo.'),
    ('ajustes-general', 'Ajustes · General',
     'Los porcentajes ideales de la regla, cómo cuenta la comida y los grupos de fijos '
     'del análisis.'),
    ('ajustes-ia', 'Ajustes · Inteligencia artificial',
     'Opcional. Sirve para proponer a qué concepto va cada línea al importar. Sin '
     'clave, todo lo demás funciona igual.'),
    ('ajustes-reglas', 'Ajustes · Reglas de clasificación',
     'Cómo se reconoce cada movimiento del banco. Se miran de arriba abajo y gana la '
     'primera que encaja, así que el orden importa tanto como el texto.'),
    ('ajustes-banco', 'Ajustes · Formato del banco',
     'Cómo se lee el archivo. Las columnas se buscan por su nombre en la cabecera, no '
     'por su posición, así que aunque el banco las mueva sigue funcionando.'),
    ('pin', 'PIN',
     'La aplicación es de la familia y vive en internet, así que pide un PIN al entrar. '
     'Es lo primero que se ve.'),
]


def parrafo(doc, texto, *, tamano=11, negrita=False, color=TINTA, espacio=6, alineado=None):
    p = doc.add_paragraph()
    if alineado is not None:
        p.alignment = alineado
    p.paragraph_format.space_after = Pt(espacio)
    t = p.add_run(texto)
    t.font.size = Pt(tamano)
    t.font.bold = negrita
    t.font.color.rgb = color
    t.font.name = 'Calibri'
    return p


def medida(ruta):
    """Ancho y alto de un PNG, leyendo su cabecera. No hace falta más."""
    with open(ruta, 'rb') as f:
        cabecera = f.read(24)
    return struct.unpack('>II', cabecera[16:24])


def imagen_si_esta(doc, ruta, ancho_cm, alto_maximo_cm=ALTO_UTIL):
    """Mete la imagen si existe. Si falta, lo dice en vez de fallar.

    Algunas pantallas son larguísimas de arriba abajo —Conceptos son sesenta
    filas— y al ancho de la página saldrían de 50 cm de alto. Cuando eso pasa se
    manda por el alto y se deja que el ancho salga como salga: más vale verla
    entera y pequeña que cortada.
    """
    if not os.path.exists(ruta):
        parrafo(doc, f'(falta la captura {os.path.basename(ruta)})', tamano=9, color=TINTA_2)
        return False

    ancho_px, alto_px = medida(ruta)
    alto_cm = ancho_cm * alto_px / ancho_px
    if alto_cm > alto_maximo_cm:
        doc.add_picture(ruta, height=Cm(alto_maximo_cm))
    else:
        doc.add_picture(ruta, width=Cm(ancho_cm))

    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_after = Pt(4)
    return True


doc = Document()

seccion = doc.sections[0]
seccion.orientation = WD_ORIENT.PORTRAIT
seccion.top_margin = Cm(1.8)
seccion.bottom_margin = Cm(1.8)
seccion.left_margin = Cm(2)
seccion.right_margin = Cm(2)
ancho_util = 21.0 - 4  # A4 menos los márgenes

# ---------------------------------------------------------------------------
# Portada
# ---------------------------------------------------------------------------
for _ in range(6):
    doc.add_paragraph()

parrafo(doc, 'gastos.', tamano=40, negrita=True, alineado=WD_ALIGN_PARAGRAPH.CENTER, espacio=2)
parrafo(doc, 'Las pantallas de la aplicación', tamano=16,
        color=TINTA_2, alineado=WD_ALIGN_PARAGRAPH.CENTER, espacio=18)
parrafo(doc, f'Generado el {date.today().strftime("%d/%m/%Y")}', tamano=10,
        color=TINTA_2, alineado=WD_ALIGN_PARAGRAPH.CENTER, espacio=2)
parrafo(doc, 'Cada pantalla, a 1120 px y a 390 px', tamano=10,
        color=TINTA_2, alineado=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ---------------------------------------------------------------------------
# Índice
# ---------------------------------------------------------------------------
parrafo(doc, 'Qué hay dentro', tamano=18, negrita=True, espacio=10)
for i, (_, titulo, _) in enumerate(PANTALLAS, start=1):
    parrafo(doc, f'{i}.  {titulo}', tamano=11, espacio=3)

doc.add_page_break()

# ---------------------------------------------------------------------------
# Una pantalla por página
# ---------------------------------------------------------------------------
for i, (clave, titulo, descripcion) in enumerate(PANTALLAS, start=1):
    parrafo(doc, titulo, tamano=18, negrita=True, espacio=4)
    parrafo(doc, descripcion, tamano=10.5, color=TINTA_2, espacio=12)

    escritorio = os.path.join(CAPTURAS, f'{clave}-escritorio.png')
    movil = os.path.join(CAPTURAS, f'{clave}-movil.png')

    if os.path.exists(escritorio):
        parrafo(doc, 'En el ordenador', tamano=9, negrita=True, color=TINTA_2, espacio=4)
        imagen_si_esta(doc, escritorio, ancho_util)

    if os.path.exists(movil):
        parrafo(doc, 'En el móvil', tamano=9, negrita=True, color=TINTA_2, espacio=4)
        # El móvil, a un tercio: si no, ocupa media página para nada.
        imagen_si_esta(doc, movil, ancho_util / 2.6, alto_maximo_cm=13)

    if i < len(PANTALLAS):
        doc.add_page_break()

doc.save(SALIDA)
print('escrito', SALIDA)
